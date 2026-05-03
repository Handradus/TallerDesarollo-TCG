from docx import Document
from docx.shared import Pt

# Data (CLP)
scenarios = {
    'Bajo': 89810,
    'Intermedio': 112310,
    'Alto': 134810,
}

# Assumptions
annual_rate = 0.08
r_month = (1+annual_rate)**(1/12)-1
months = 60
factor_pv = (1 - (1+r_month)**(-months)) / r_month

# Create document
doc = Document()

doc.add_heading('VAN_CLP — Valor Actual Neto (CLP)', level=1)

doc.add_heading('Supuestos', level=2)
doc.add_paragraph(f'Tasa anual: {annual_rate*100:.2f}% (tasa mensual ≈ {r_month:.6f})')
doc.add_paragraph(f'Horizonte: {months} meses')
doc.add_paragraph(f'Factor VAN (mensual): {factor_pv:.2f} (PV factor = (1 - (1+r)^-{months})/r)')

doc.add_heading('Resultados (CLP)', level=2)
for name, monthly in scenarios.items():
    pv = int(round(monthly * factor_pv))
    p = doc.add_paragraph()
    p.add_run(f'- Escenario {name} ({monthly:,} CLP/mes): ').bold = True
    p.add_run(f'VAN_CLP ≈ {pv:,} CLP')

# Detailed table
doc.add_heading('Detalle de cálculos', level=2)
doc.add_paragraph('Fórmulas:')
doc.add_paragraph('  r = (1 + tasa_anual)^(1/12) - 1')
doc.add_paragraph('  Factor PV = (1 - (1+r)^-N) / r')
doc.add_paragraph('  VAN_CLP = mensualidad * Factor PV')

doc.add_heading('Valores numéricos', level=3)
doc.add_paragraph(f'Tasa anual: {annual_rate*100:.2f}%')
doc.add_paragraph(f'Tasa mensual r: {r_month:.6f}')
doc.add_paragraph(f'Factor PV (60 meses): {factor_pv:.6f}')

doc.add_heading('Supuestos de costos mensuales', level=3)
doc.add_paragraph('Escenario Bajo: 89.810 CLP/mes')
doc.add_paragraph('Escenario Intermedio: 112.310 CLP/mes')
doc.add_paragraph('Escenario Alto: 134.810 CLP/mes')

# Save
out_path = 'docs/VAN_CLP_calculado.docx'
doc.save(out_path)
print('WROTE', out_path)
